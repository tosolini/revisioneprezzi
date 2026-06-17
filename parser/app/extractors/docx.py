from docx import Document

from app.extractors.base import BaseExtractor


class DocxExtractor(BaseExtractor):
    def extract_text(self, file_path: str) -> str:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables.append(" | ".join(cells))
        return "\n".join(paragraphs + tables)
